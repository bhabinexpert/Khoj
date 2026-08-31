"""Turn an uploaded PDF/DOCX into plain text.

Kept separate from parsing so the two concerns can fail independently: a
password-protected PDF is an extraction problem, a CV with no "Skills" heading
is a parsing problem, and the API reports them differently.
"""

from __future__ import annotations

import io
import logging
import re
from dataclasses import dataclass, field

from app.config import settings

logger = logging.getLogger(__name__)


class ExtractionError(RuntimeError):
    """Raised when no text can be recovered from the upload."""


@dataclass
class ExtractedText:
    text: str
    file_type: str  # "pdf" | "docx" | "unknown"
    pages: int = 0
    warnings: list[str] = field(default_factory=list)


def detect_file_type(filename: str, content: bytes) -> str:
    """Sniff the real type from magic bytes, falling back to the extension.

    Browsers lie about MIME types for .docx often enough that the bytes are the
    more reliable signal.
    """
    if content[:5] == b"%PDF-":
        return "pdf"
    # .docx is a zip; so is .xlsx, but the caller has already filtered by name.
    if content[:2] == b"PK" and (filename or "").lower().endswith(".docx"):
        return "docx"
    lowered = (filename or "").lower()
    if lowered.endswith(".pdf"):
        return "pdf"
    if lowered.endswith((".docx", ".doc")):
        return "docx"
    return "unknown"


def _clean(text: str) -> str:
    """Normalise the whitespace PDF extraction inevitably mangles."""
    text = text.replace("\r\n", "\n").replace("\r", "\n").replace("\xa0", " ")
    # Ligatures and dashes that break keyword matching.
    for needle, replacement in (("ﬁ", "fi"), ("ﬂ", "fl"), ("–", "-"), ("—", "-"), ("•", "•")):
        text = text.replace(needle, replacement)
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.split("\n")]
    out: list[str] = []
    for line in lines:
        if not line and out and not out[-1]:
            continue
        out.append(line)
    return "\n".join(out).strip()


def extract_pdf(content: bytes) -> ExtractedText:
    import pdfplumber

    warnings: list[str] = []
    chunks: list[str] = []

    try:
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            total_pages = len(pdf.pages)
            if total_pages > settings.max_pdf_pages:
                warnings.append(
                    f"Only the first {settings.max_pdf_pages} of {total_pages} pages were read."
                )
            for page in pdf.pages[: settings.max_pdf_pages]:
                chunks.append(page.extract_text() or "")
    except Exception as exc:  # noqa: BLE001 - pdfplumber raises a wide variety
        raise ExtractionError(
            "Could not read that PDF. It may be corrupt or password-protected."
        ) from exc

    text = _clean("\n".join(chunks))
    if not text:
        raise ExtractionError(
            "No text found in that PDF — it looks like a scanned image. "
            "Please upload a text-based PDF or a DOCX file."
        )
    return ExtractedText(text=text, file_type="pdf", pages=total_pages, warnings=warnings)


def extract_docx(content: bytes) -> ExtractedText:
    import docx

    try:
        document = docx.Document(io.BytesIO(content))
    except Exception as exc:  # noqa: BLE001
        raise ExtractionError("Could not read that DOCX file. It may be corrupt or a legacy .doc.") from exc

    parts = [paragraph.text for paragraph in document.paragraphs]
    # Plenty of CVs lay everything out in tables; paragraphs alone would miss it.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    text = _clean("\n".join(parts))
    if not text:
        raise ExtractionError("That DOCX file has no readable text in it.")
    return ExtractedText(text=text, file_type="docx", pages=0)


def extract_text(filename: str, content: bytes) -> ExtractedText:
    """Extract text from ``content``, dispatching on the sniffed file type."""
    if not content:
        raise ExtractionError("The uploaded file is empty.")
    if len(content) > settings.max_upload_bytes:
        raise ExtractionError(
            f"File is too large ({len(content) / 1_048_576:.1f} MB). "
            f"Maximum is {settings.max_upload_bytes / 1_048_576:.0f} MB."
        )

    file_type = detect_file_type(filename, content)
    if file_type == "pdf":
        result = extract_pdf(content)
    elif file_type == "docx":
        result = extract_docx(content)
    else:
        raise ExtractionError("Unsupported file type. Please upload a PDF or DOCX.")

    if len(result.text) > settings.max_text_chars:
        result.warnings.append("Document was very long; only the first portion was analysed.")
        result.text = result.text[: settings.max_text_chars]

    logger.info("extracted %s chars from %s (%s)", len(result.text), filename, file_type)
    return result
