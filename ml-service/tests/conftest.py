"""Shared fixtures.

``EMBEDDINGS_ENABLED=false`` is set *before* ``app.config`` is imported so the
suite exercises the deterministic lexical matcher and never downloads 2.5 GB of
torch in CI. ``test_match_score.py`` has a separate, explicitly-skipped test for
the semantic path.
"""

from __future__ import annotations

import io
import os

os.environ.setdefault("EMBEDDINGS_ENABLED", "false")

import pytest  # noqa: E402
from fastapi.testclient import TestClient  # noqa: E402

from app.main import app  # noqa: E402

SAMPLE_CV = """Sabin Shrestha
Kathmandu, Nepal | +977-9801234567 | sabin.shrestha@example.com

PROFESSIONAL SUMMARY
Frontend developer building React applications for fintech teams.

TECHNICAL SKILLS
Languages: JavaScript, TypeScript, Python
Frontend: ReactJS, Redux Toolkit, Tailwind CSS, HTML5, CSS3
Backend: Node.js, Express.js, MongoDB, REST API
Tools: Git, Docker, Figma
Soft skills: Communication, Teamwork, Problem Solving

WORK EXPERIENCE
Frontend Developer at Leapfrog Technology Pvt. Ltd.
Jan 2022 - Dec 2024
- Built dashboards in React and Redux, cutting load time by 40%.

Web Development Intern
Yarsa Labs | Jun 2021 - Dec 2021
- Maintained jQuery and PHP legacy pages.

EDUCATION
BSc. CSIT in Computer Science, Tribhuvan University, 2021
+2 Science, Kathmandu Model College, 2017

CERTIFICATIONS
AWS Certified Cloud Practitioner (2023)
"""


@pytest.fixture(scope="session")
def client() -> TestClient:
    return TestClient(app)


@pytest.fixture(scope="session")
def sample_pdf() -> bytes:
    """A real, text-based PDF built with reportlab."""
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    y = height - 50
    for line in SAMPLE_CV.split("\n"):
        if y < 50:
            pdf.showPage()
            y = height - 50
        pdf.setFont("Helvetica", 10)
        pdf.drawString(45, y, line[:110])
        y -= 14
    pdf.save()
    return buffer.getvalue()


@pytest.fixture(scope="session")
def sample_docx() -> bytes:
    """A DOCX that puts the skills in a table, like many real CVs do."""
    import docx

    document = docx.Document()
    for line in SAMPLE_CV.split("\n"):
        if line.startswith("Languages:") or line.startswith("Frontend:"):
            continue
        document.add_paragraph(line)
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Languages"
    table.cell(0, 1).text = "JavaScript, TypeScript, Python"
    table.cell(1, 0).text = "Frontend"
    table.cell(1, 1).text = "ReactJS, Redux Toolkit, Tailwind CSS"
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


@pytest.fixture(autouse=True)
def _fresh_matcher():
    """Never let one test's matcher leak into another."""
    from app.matching.embedder import reset_matcher

    reset_matcher()
    yield
    reset_matcher()
